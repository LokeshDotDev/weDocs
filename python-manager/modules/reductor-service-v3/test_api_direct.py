#!/usr/bin/env python3
"""
Test the FULL backend flow:
1. Create a test DOCX with student info
2. Call the reductor API endpoint
3. Verify the output
"""

import requests
import json
import tempfile
import os
from pathlib import Path
from docx import Document

# Create test document with student info
def create_test_docx():
    doc = Document()
    doc.add_paragraph("NAME: SHANMUGAPRIYA SIVAKUMAR")
    doc.add_paragraph("ROLL NUMBER: 251450500104")
    doc.add_paragraph("PROGRAM: MASTER OF BUSINESS ADMINISTRATION(MBA)")
    doc.add_paragraph("SEMESTER:1")
    doc.add_paragraph("COURSE CODE & NAME: DMBA114- BUSINESS COMMUNICATION")
    doc.add_paragraph("")
    doc.add_paragraph("Assignment Content")
    doc.add_paragraph("This is the assignment submitted by SHANMUGAPRIYA SIVAKUMAR (Roll: 251450500104)")
    
    # Save and return path
    tmpfile = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmpfile.name)
    tmpfile.close()
    return tmpfile.name

def verify_docx(filepath):
    """Read and display DOCX content"""
    doc = Document(filepath)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)

def test_api():
    """Test the reductor API directly"""
    
    print("=" * 80)
    print("TESTING REDUCTOR SERVICE V3 API")
    print("=" * 80)
    
    # Create test document
    print("\n1. Creating test document...")
    test_file = create_test_docx()
    print(f"   ✓ Created: {test_file}")
    
    # Read original content
    print("\n2. Original document content:")
    print("-" * 80)
    original_content = verify_docx(test_file)
    print(original_content)
    print("-" * 80)
    
    # Create output path
    output_file = test_file.replace('.docx', '_output.docx')
    
    # Call the reductor service directly
    print(f"\n3. Calling reductor-service-v3 API...")
    print(f"   Input: {test_file}")
    print(f"   Output: {output_file}")
    
    try:
        response = requests.post(
            'http://localhost:5018/redact/document',
            json={
                'input_file_path': test_file,
                'output_file_path': output_file,
                'file_format': 'docx',
                'remove_name': True,
                'remove_roll_no': True
            },
            timeout=30
        )
        
        print(f"   Status: {response.status_code}")
        result = response.json()
        print(f"   Response: {json.dumps(result, indent=2)}")
        
        if response.status_code != 200:
            print("   ❌ API call failed!")
            return False
        
        # Check if output file exists
        if not os.path.exists(output_file):
            print(f"   ❌ Output file not created: {output_file}")
            return False
        
        print(f"   ✓ Output file created: {output_file}")
        
        # Read redacted content
        print("\n4. Redacted document content:")
        print("-" * 80)
        redacted_content = verify_docx(output_file)
        print(redacted_content)
        print("-" * 80)
        
        # Verify redaction
        print("\n5. Verification:")
        
        name_found = "SHANMUGAPRIYA SIVAKUMAR" in redacted_content
        roll_found = "251450500104" in redacted_content
        redacted_found = "[REDACTED]" in redacted_content
        
        print(f"   NAME present: {name_found} (should be False) {'❌' if name_found else '✓'}")
        print(f"   ROLL NUMBER present: {roll_found} (should be False) {'❌' if roll_found else '✓'}")
        print(f"   [REDACTED] found: {redacted_found} (should be True) {'✓' if redacted_found else '❌'}")
        
        # Cleanup
        os.unlink(test_file)
        os.unlink(output_file)
        
        if not name_found and not roll_found and redacted_found:
            print("\n" + "=" * 80)
            print("✓✓✓ API TEST PASSED ✓✓✓")
            print("=" * 80)
            return True
        else:
            print("\n" + "=" * 80)
            print("❌ API TEST FAILED")
            print("=" * 80)
            return False
            
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    import sys
    success = test_api()
    sys.exit(0 if success else 1)
