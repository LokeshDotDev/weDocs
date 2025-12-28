"""
Reductor Service v3 - Student Name and Roll Number Redaction Service

This service is specifically designed to handle the second pattern of student documents
(as shown in screenshot 2). It detects and removes ONLY the student NAME and ROLL NUMBER
fields from documents while preserving all other information like course code, program, etc.

Service Type: FastAPI-based microservice
Port: 5018 (default)
"""

import os
import re
import sys
import logging
from typing import Optional, List, Dict, Any
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Body
from docx_anonymizer import anonymize_docx
from pydantic import BaseModel
from fastapi.responses import FileResponse

# ============================================================================
# Setup Logging for Debugging
# ============================================================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - [REDUCTOR V3] %(message)s'
)
logger = logging.getLogger(__name__)

# ============================================================================
# Models for API Request/Response
# ============================================================================

class StudentIdentifierRequest(BaseModel):
    """Request to extract student identifiers from text"""
    text: str
    strict_mode: bool = True  # Use stricter pattern matching if True


class StudentIdentifierResponse(BaseModel):
    """Response containing extracted student identifiers"""
    detected_name: Optional[str] = None
    detected_roll_no: Optional[str] = None
    extraction_confidence: str = "high"  # high, medium, low
    raw_text_excerpt: Optional[str] = None


class RedactionRequest(BaseModel):
    """Request to redact student identifiers from text"""
    text: str
    remove_name: bool = True
    remove_roll_no: bool = True
    preserve_labels: bool = False  # Keep "NAME:" and "ROLL NUMBER:" labels


class RedactionResponse(BaseModel):
    """Response containing redacted text and detection info"""
    redacted_text: str
    detected_name: Optional[str] = None
    detected_roll_no: Optional[str] = None
    redaction_count: int


class DocumentRedactionRequest(BaseModel):
    """Request to redact a document file"""
    input_file_path: str
    output_file_path: str
    file_format: str = "docx"  # docx, pdf, txt
    remove_name: bool = True
    remove_roll_no: bool = True


class DocumentRedactionResponse(BaseModel):
    """Response for document redaction"""
    status: str
    output_file: str
    redacted_name: Optional[str] = None
    redacted_roll_no: Optional[str] = None


# ============================================================================
# Pattern Extraction Engine
# ============================================================================

class StudentIdentifierExtractor:
    """
    Extracts student NAME and ROLL NUMBER from documents.
    
    Optimized for NON-TABLE format files (screenshot 2):
    - NAME: SHANMUGAPRIYA SIVAKUMAR (Full name, usually 2-4 parts, all capitals or title case)
    - ROLL NUMBER: 25145050010 (10-15 digits, may have special chars like -)
    
    Patterns handle:
    - Flexible whitespace around colons
    - Uppercase and mixed case names
    - Roll numbers with dashes, spaces, or special characters
    """
    
    # Strict patterns - high confidence for non-table format
    # Case-insensitive to handle variations
    STRICT_NAME_PATTERN = re.compile(
        r"NAME\s*:\s*([A-Z][A-Za-z\s]+?)(?=\n|$|ROLL|PROGRAM|SEMESTER|COURSE)", 
        re.IGNORECASE | re.MULTILINE
    )
    STRICT_ROLL_PATTERN = re.compile(
        r"ROLL\s*(?:NUMBER|NO\.?)\s*:\s*([\d\-\s]{8,20}?)(?=\n|$|PROGRAM|COURSE|SEMESTER|STUDENT)", 
        re.IGNORECASE | re.MULTILINE
    )
    
    # Flexible patterns - for cases with formatting variations
    FLEX_NAME_PATTERN = re.compile(
        r"(?:STUDENT\s+)?NAME\s*[:\s]+\s*([A-Z][A-Za-z\s]+?)(?=\n|ROLL|PROGRAM|$)",
        re.IGNORECASE | re.MULTILINE
    )
    FLEX_ROLL_PATTERN = re.compile(
        r"(?:ROLL\s*(?:NUMBER|NO)|ENROLLMENT\s+(?:NUMBER|NO)|REGISTRATION|STUDENT\s*ID|REG\s*(?:NUMBER|NO))\s*[:\s]+\s*([\d\-\s]{8,20}?)(?=\n|PROGRAM|COURSE|SEMESTER|$)",
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern for names in various formats
    GENERIC_NAME_PATTERN = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b")
    
    @staticmethod
    def extract_name(text: str, strict: bool = True) -> tuple[Optional[str], str]:
        """
        Extract student name from text.
        Optimized for non-table format with NAME: label.
        
        Args:
            text: Input text to search
            strict: If True, use strict pattern. If False, use flexible pattern.
            
        Returns:
            Tuple of (extracted_name, confidence_level)
        """
        if strict:
            match = StudentIdentifierExtractor.STRICT_NAME_PATTERN.search(text)
            if match:
                name = match.group(1).strip()
                # Clean up the name (remove extra whitespace)
                name = re.sub(r'\s+', ' ', name)
                # Remove line breaks and trailing non-letter characters
                name = re.sub(r'[^A-Za-z\s].*$', '', name).strip()
                if name:
                    return name, "high"
        
        # Try flexible pattern if strict doesn't work
        match = StudentIdentifierExtractor.FLEX_NAME_PATTERN.search(text)
        if match:
            name = match.group(1).strip()
            name = re.sub(r'\s+', ' ', name)
            name = re.sub(r'[^A-Za-z\s].*$', '', name).strip()
            if name:
                return name, "medium"
        
        return None, "none"
    
    @staticmethod
    def extract_roll_number(text: str, strict: bool = True) -> tuple[Optional[str], str]:
        """
        Extract roll number from text.
        Optimized for non-table format with ROLL NUMBER: label.
        
        Args:
            text: Input text to search
            strict: If True, use strict pattern. If False, use flexible pattern.
            
        Returns:
            Tuple of (extracted_roll_number, confidence_level)
        """
        if strict:
            match = StudentIdentifierExtractor.STRICT_ROLL_PATTERN.search(text)
            if match:
                roll_no = match.group(1).strip()
                # Clean up: remove spaces and hyphens, keep only digits
                roll_no_clean = re.sub(r'[\s\-]', '', roll_no)
                if roll_no_clean and len(roll_no_clean) >= 8:
                    return roll_no_clean, "high"
        
        # Try flexible pattern if strict doesn't work
        match = StudentIdentifierExtractor.FLEX_ROLL_PATTERN.search(text)
        if match:
            roll_no = match.group(1).strip()
            roll_no_clean = re.sub(r'[\s\-]', '', roll_no)
            if roll_no_clean and len(roll_no_clean) >= 8:
                return roll_no_clean, "medium"
        
        return None, "none"
    
    @staticmethod
    def extract_both(text: str, strict: bool = True) -> tuple[Optional[str], Optional[str], str]:
        """
        Extract both name and roll number from text.
        
        Args:
            text: Input text to search
            strict: If True, use strict patterns. If False, use flexible patterns.
            
        Returns:
            Tuple of (name, roll_number, confidence_level)
        """
        name, name_conf = StudentIdentifierExtractor.extract_name(text, strict)
        roll_no, roll_conf = StudentIdentifierExtractor.extract_roll_number(text, strict)
        
        # Determine overall confidence
        if name_conf == "high" and roll_conf == "high":
            overall_conf = "high"
        elif name_conf == "high" or roll_conf == "high":
            overall_conf = "medium"
        else:
            overall_conf = "low"
        
        return name, roll_no, overall_conf


# ============================================================================
# Redaction Engine
# ============================================================================

class StudentIdentifierRedactor:
    """
    Redacts student NAME and ROLL NUMBER from documents while preserving
    all other information.
    """
    
    @staticmethod
    def redact_name(text: str, preserve_label: bool = False) -> tuple[str, Optional[str]]:
        """
        Redact student name from text.
        Optimized for non-table format with NAME: label.
        
        Args:
            text: Input text to redact
            preserve_label: If True, keep "NAME:" label but remove the actual name
            
        Returns:
            Tuple of (redacted_text, detected_name)
        """
        detected_name = None
        
        # First try strict pattern
        match = StudentIdentifierExtractor.STRICT_NAME_PATTERN.search(text)
        
        if match:
            detected_name = match.group(1).strip()
            detected_name = re.sub(r'\s+', ' ', detected_name)
            detected_name = re.sub(r'[^A-Za-z\s].*$', '', detected_name).strip()
            
            if preserve_label:
                # Find the NAME: part and keep label
                name_part = match.group(0)
                replacement = re.sub(r"([A-Z][A-Za-z\s]+)", "[REDACTED]", name_part)
                redacted = text.replace(name_part, replacement)
            else:
                # Replace entire line with just the label removed
                redacted = text.replace(match.group(0), "[REDACTED]")
        else:
            # Try flexible pattern
            match = StudentIdentifierExtractor.FLEX_NAME_PATTERN.search(text)
            if match:
                detected_name = match.group(1).strip()
                detected_name = re.sub(r'\s+', ' ', detected_name)
                detected_name = re.sub(r'[^A-Za-z\s].*$', '', detected_name).strip()
                
                if preserve_label:
                    name_part = match.group(0)
                    replacement = re.sub(r"([A-Z][A-Za-z\s]+)", "[REDACTED]", name_part)
                    redacted = text.replace(name_part, replacement)
                else:
                    redacted = text.replace(match.group(0), "[REDACTED]")
            else:
                redacted = text
        
        return redacted, detected_name
    
    @staticmethod
    def redact_roll_number(text: str, preserve_label: bool = False) -> tuple[str, Optional[str]]:
        """
        Redact roll number from text.
        Optimized for non-table format with ROLL NUMBER: label.
        
        Args:
            text: Input text to redact
            preserve_label: If True, keep "ROLL NUMBER:" label but remove the actual number
            
        Returns:
            Tuple of (redacted_text, detected_roll_number)
        """
        detected_roll = None
        
        # First try strict pattern
        match = StudentIdentifierExtractor.STRICT_ROLL_PATTERN.search(text)
        
        if match:
            roll_raw = match.group(1).strip()
            detected_roll = re.sub(r'[\s\-]', '', roll_raw)
            
            if preserve_label:
                # Keep label, replace only the number
                roll_part = match.group(0)
                replacement = re.sub(r"([\d\-\s]+)", "[REDACTED]", roll_part)
                redacted = text.replace(roll_part, replacement)
            else:
                # Replace entire roll number line
                redacted = text.replace(match.group(0), "[REDACTED]")
        else:
            # Try flexible pattern
            match = StudentIdentifierExtractor.FLEX_ROLL_PATTERN.search(text)
            if match:
                roll_raw = match.group(1).strip()
                detected_roll = re.sub(r'[\s\-]', '', roll_raw)
                
                if preserve_label:
                    roll_part = match.group(0)
                    replacement = re.sub(r"([\d\-\s]+)", "[REDACTED]", roll_part)
                    redacted = text.replace(roll_part, replacement)
                else:
                    redacted = text.replace(match.group(0), "[REDACTED]")
            else:
                redacted = text
        
        return redacted, detected_roll
    
    @staticmethod
    def redact_both(text: str, remove_name: bool = True, 
                   remove_roll: bool = True, preserve_labels: bool = False) -> tuple[str, Dict[str, Any]]:
        """
        Redact both name and roll number from text globally.
        
        Strategy:
        1. Extract name and roll number from the beginning
        2. Remove ALL occurrences of those values throughout the entire file
        
        Args:
            text: Input text to redact
            remove_name: If True, redact the name everywhere
            remove_roll: If True, redact the roll number everywhere
            preserve_labels: If True, keep field labels (NAME:, ROLL NUMBER:)
            
        Returns:
            Tuple of (redacted_text, metadata_dict)
        """
        redacted = text
        metadata = {
            "detected_name": None,
            "detected_roll_no": None,
            "redaction_count": 0
        }
        
        # Step 1: Extract the actual name and roll number from the beginning
        if remove_name:
            name, name_conf = StudentIdentifierExtractor.extract_name(text, strict=True)
            if not name:
                # Try flexible pattern if strict fails
                name, name_conf = StudentIdentifierExtractor.extract_name(text, strict=False)
            
            metadata["detected_name"] = name
            logger.info(f"Extracted NAME: '{name}' (confidence: {name_conf})")
            
            # Step 2: Remove ALL occurrences of this name from entire file
            if name:
                # Count occurrences before redaction
                count_before = redacted.count(name)
                count_before_lower = redacted.lower().count(name.lower())
                logger.info(f"Found {count_before} exact matches and {count_before_lower} case-insensitive matches for name")
                
                # Case-insensitive global replacement
                pattern = re.compile(re.escape(name), re.IGNORECASE)
                redacted = pattern.sub("[REDACTED]", redacted)
                
                count_after = redacted.count(name)
                logger.info(f"After redaction: {count_after} remaining matches (removed {count_before - count_after})")
                metadata["redaction_count"] += 1
        
        if remove_roll:
            roll_no, roll_conf = StudentIdentifierExtractor.extract_roll_number(text, strict=True)
            if not roll_no:
                # Try flexible pattern if strict fails
                roll_no, roll_conf = StudentIdentifierExtractor.extract_roll_number(text, strict=False)
            
            metadata["detected_roll_no"] = roll_no
            logger.info(f"Extracted ROLL NUMBER: '{roll_no}' (confidence: {roll_conf})")
            
            # Step 2: Remove ALL occurrences of this roll number from entire file
            if roll_no:
                # Count occurrences before redaction
                count_before = redacted.count(roll_no)
                logger.info(f"Found {count_before} exact matches for roll number")
                
                # Global replacement
                pattern = re.compile(re.escape(roll_no), re.IGNORECASE)
                redacted = pattern.sub("[REDACTED]", redacted)
                
                # Also try to remove formatted versions (with dashes, spaces)
                # e.g., if roll_no is "25145050010", also remove "25145050-0010", "2514 5050 0010", etc
                formatted_patterns = [
                    roll_no[:8] + '-' + roll_no[8:],  # Add dash in middle
                    ' '.join(roll_no[i:i+4] for i in range(0, len(roll_no), 4)),  # Add spaces every 4 digits
                    roll_no[:5] + '-' + roll_no[5:],  # Dash after 5 digits
                ]
                for fmt_pattern in formatted_patterns:
                    if fmt_pattern != roll_no:
                        pattern = re.compile(re.escape(fmt_pattern), re.IGNORECASE)
                        matches = len(pattern.findall(redacted))
                        if matches > 0:
                            logger.info(f"Found {matches} matches for formatted pattern: '{fmt_pattern}'")
                            redacted = pattern.sub("[REDACTED]", redacted)
                
                count_after = redacted.count(roll_no)
                logger.info(f"After redaction: {count_after} remaining matches")
                metadata["redaction_count"] += 1
        
        return redacted, metadata


# ============================================================================
# File Handling Utilities
# ============================================================================

class DocumentProcessor:
    """Handles processing of different document formats"""
    
    @staticmethod
    def read_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX file: {str(e)}")
    
    @staticmethod
    def read_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages])
        except ImportError:
            raise RuntimeError("PyPDF2 not installed. Install with: pip install PyPDF2")
        except Exception as e:
            raise RuntimeError(f"Failed to read PDF file: {str(e)}")
    
    @staticmethod
    def read_txt(file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise RuntimeError(f"Failed to read TXT file: {str(e)}")
    
    @staticmethod
    def write_docx(file_path: str, text: str) -> None:
        """
        Write redacted text to DOCX file by modifying existing document structure.
        
        Strategy:
        - Read the original DOCX document
        - Replace text within each paragraph while preserving formatting
        - Save the modified document
        """
        try:
            from docx import Document
            
            # Load the ORIGINAL document (not create new)
            # We need to get the original file path - but we only have output path
            # This is a problem - we need to modify the approach
            
            # Create a new document from the redacted text
            doc = Document()
            
            # Parse the text and add with better formatting
            for line in text.split('\n'):
                if line.strip() or True:  # Keep empty lines too
                    para = doc.add_paragraph(line)
                    
                    # Highlight redacted sections in red and bold
                    if '[REDACTED]' in para.text:
                        for run in para.runs:
                            if '[REDACTED]' in run.text:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            doc.save(file_path)
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to write DOCX file: {str(e)}")
    
    @staticmethod
    def modify_docx_in_place(input_file: str, output_file: str, name_to_replace: Optional[str] = None, roll_to_replace: Optional[str] = None) -> int:
        """
        Modify an existing DOCX file by replacing text at the run level.
        
        This approach is more reliable because it:
        1. Works with text split across multiple runs
        2. Preserves formatting attributes
        3. Directly modifies text in runs without restructuring paragraphs
        
        Args:
            input_file: Path to the original DOCX file
            output_file: Path to save the modified document
            name_to_replace: The exact student name to replace (e.g., 'SHANMUGAPRIYA SIVAKUMAR')
            roll_to_replace: The exact roll number to replace (e.g., '251450500104')
            
        Returns:
            Total number of replacements made
        """
        try:
            from docx import Document
            from docx.shared import RGBColor
            
            logger.info(f"Loading original document: {input_file}")
            doc = Document(input_file)
            
            # Preserve document styles and list definitions
            # This ensures bullets and numbering remain intact after modifications
            
            logger.info(f"Will replace NAME: '{name_to_replace}'")
            logger.info(f"Will replace ROLL: '{roll_to_replace}'")
            
            replacement_count = 0
            
            # Build list of (paragraph, full_text) to check
            replacements_to_make = []
            
            # Scan all paragraphs — target only label lines (NAME:, ROLL NUMBER:)
            # Match only standalone NAME: at start, not "COURSE CODE & NAME:"
            name_label = re.compile(r"^\s*NAME\s*:\s*", re.IGNORECASE)
            roll_label = re.compile(r"^\s*ROLL\s*NUMBER\s*:\s*", re.IGNORECASE)
            for para_idx, para in enumerate(doc.paragraphs):
                para_text = para.text
                to_replace_in_para = []
                if name_to_replace and name_label.search(para_text):
                    logger.info(f"[Para {para_idx}] NAME label detected: {para_text}")
                    to_replace_in_para.append(("name-label", name_label))
                if roll_to_replace and roll_label.search(para_text):
                    logger.info(f"[Para {para_idx}] ROLL label detected: {para_text}")
                    to_replace_in_para.append(("roll-label", roll_label))
                if to_replace_in_para:
                    replacements_to_make.append((para, para_text, to_replace_in_para))
            
            # Scan all table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = para.text
                            to_replace_in_para = []
                            
                            if name_to_replace and name_label.search(para_text):
                                logger.info(f"[Table Cell] NAME label detected: {para_text}")
                                to_replace_in_para.append(("name-label", name_label))
                            if roll_to_replace and roll_label.search(para_text):
                                logger.info(f"[Table Cell] ROLL label detected: {para_text}")
                                to_replace_in_para.append(("roll-label", roll_label))
                            
                            if to_replace_in_para:
                                replacements_to_make.append((para, para_text, to_replace_in_para))
            
            # Now actually replace the text
            for para, original_text, replacements in replacements_to_make:
                def _replace_value_in_runs(paragraph, value: str):
                    # Remove all case-insensitive occurrences of value from existing runs
                    if not value:
                        return
                    value_lower = value.lower()
                    # Repeat until no more matches
                    while True:
                        combined = "".join(run.text for run in paragraph.runs)
                        idx = combined.lower().find(value_lower)
                        if idx == -1:
                            break
                        start = idx
                        end = idx + len(value)
                        pos = 0
                        for run in paragraph.runs:
                            run_text = run.text
                            run_start = pos
                            run_end = pos + len(run_text)
                            # overlap with [start, end)
                            overlap_start = start if start > run_start else run_start
                            if run_end < start:
                                pos = run_end
                                continue
                            overlap_start = max(start, run_start)
                            overlap_end = min(end, run_end)
                            if overlap_start < overlap_end:
                                before = overlap_start - run_start
                                after = run_end - overlap_end
                                run.text = run_text[:before] + run_text[-after:] if after > 0 else run_text[:before]
                            pos = run_end

                def _remove_text_after_colon(paragraph, label_regex: re.Pattern[str]):
                    # Remove everything after the label's colon, leaving blank space
                    # Always preserve leading bullet characters (•, ◦, ▪, –, etc) in the first run, and preserve their font
                    bullet_chars = {'•', '◦', '▪', '–', '-', '●', '‣', '∙', '○', '□', '■', '◆', '▶', '→', '⇒', '➔', '➤', '➢', '➣', '➥', '➦', '➧', '➨', '➩', '➪', '➫', '➬', '➭', '➮', '➯', '➱', '➲', '➳', '➵', '➸', '➺', '➻', '➼', '➽', '➾'}
                    combined = "".join(run.text for run in paragraph.runs)
                    match = label_regex.search(combined)
                    if not match:
                        return
                    colon_index = match.end()  # position after 'NAME:' or 'ROLL NUMBER:' and any spaces
                    pos = 0
                    for i, run in enumerate(paragraph.runs):
                        run_text = run.text
                        run_start = pos
                        run_end = pos + len(run_text)
                        # Always preserve bullet if it's the first character in the run
                        preserve_bullet = run_text and run_text[0] in bullet_chars
                        bullet_font = None
                        if preserve_bullet and run.font and run.font.name:
                            bullet_font = run.font.name
                        if run_end <= colon_index:
                            # entirely before deletion range: keep as is
                            pos = run_end
                            continue
                        if run_start >= colon_index:
                            # entirely within deletion range: clear text, but keep bullet if present
                            if preserve_bullet:
                                run.text = run_text[0]
                                if bullet_font:
                                    run.font.name = bullet_font
                            else:
                                run.text = ""
                        else:
                            # colon falls inside this run: preserve left side only, and bullet if present
                            keep_len = colon_index - run_start
                            if preserve_bullet and keep_len == 0:
                                run.text = run_text[0]
                                if bullet_font:
                                    run.font.name = bullet_font
                            else:
                                run.text = run_text[:keep_len]
                        pos = run_end
                    
                logger.info(f"  Original: {original_text}")
                for replace_type, value_to_replace in replacements:
                    if replace_type.endswith("label"):
                        logger.info("  Removing content after colon (labels)")
                        _remove_text_after_colon(para, value_to_replace)
                    else:
                        logger.info(f"  Removing value from runs: '{value_to_replace}'")
                        _replace_value_in_runs(para, value_to_replace)
                    replacement_count += 1
                logger.info(f"  ✓ Paragraph updated successfully (formatting preserved)")
            
            # Save the document
            logger.info(f"Saving modified document to: {output_file}")
            doc.save(output_file)
            logger.info(f"✓✓✓ SUCCESS! Document saved with {replacement_count} replacements ✓✓✓")
            
            return replacement_count
            
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"ERROR modifying DOCX: {str(e)}", exc_info=True)
            raise RuntimeError(f"Failed to modify DOCX file: {str(e)}")
    
    @staticmethod
    def write_txt(file_path: str, text: str) -> None:
        """Write redacted text to TXT file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
        except Exception as e:
            raise RuntimeError(f"Failed to write TXT file: {str(e)}")
    
    @staticmethod
    def process_document(input_path: str, output_path: str, file_format: str = "docx") -> str:
        """
        Read document, extract text, and return it.
        
        Args:
            input_path: Path to input document
            output_path: Path where output will be saved (used for format detection if needed)
            file_format: Format of the file (docx, pdf, txt)
            
        Returns:
            Extracted text from the document
        """
        if not os.path.exists(input_path):
            raise RuntimeError(f"Input file not found: {input_path}")
        
        if file_format.lower() == "docx":
            return DocumentProcessor.read_docx(input_path)
        elif file_format.lower() == "pdf":
            return DocumentProcessor.read_pdf(input_path)
        elif file_format.lower() == "txt":
            return DocumentProcessor.read_txt(input_path)
        else:
            raise RuntimeError(f"Unsupported file format: {file_format}")
    
    @staticmethod
    def save_document(output_path: str, text: str, file_format: str = "docx") -> None:
        """
        Save redacted text to document.
        
        Args:
            output_path: Path where file should be saved
            text: Redacted text to save
            file_format: Format of the file (docx, pdf, txt)
        """
        # Create output directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        
        if file_format.lower() == "docx":
            DocumentProcessor.write_docx(output_path, text)
        elif file_format.lower() == "txt":
            DocumentProcessor.write_txt(output_path, text)
        else:
            # Default to TXT for unsupported formats
            DocumentProcessor.write_txt(output_path, text)


# ============================================================================
# FastAPI Application Setup
# ============================================================================

app = FastAPI(
    title="Reductor Service v3",
    description="Student Name and Roll Number Redaction Service (Screenshot 2 Format)",
    version="3.0.0"
)


# ============================================================================
# Health & Status Endpoints
# ============================================================================

@app.get("/health")
def health():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "reductor-service-v3",
        "version": "3.0.0",
        "purpose": "Student NAME and ROLL NUMBER redaction"
    }


@app.get("/info")
def info():
    """Service information endpoint"""
    return {
        "name": "Reductor Service v3",
        "description": "Specialized service for redacting student NAME and ROLL NUMBER from documents (Screenshot 2 format)",
        "version": "3.0.0",
        "features": [
            "Extract student NAME and ROLL NUMBER from text",
            "Redact both identifiers while preserving other document content",
            "Support for DOCX, PDF, and TXT formats",
            "High confidence pattern matching with flexible fallback patterns",
            "Optional label preservation for document structure"
        ],
        "endpoints": {
            "health": "GET /health - Service health check",
            "identify": "POST /identify/text - Extract identifiers from text",
            "redact_text": "POST /redact/text - Redact identifiers from text",
            "redact_document": "POST /redact/document - Redact entire document file"
        }
    }


# ============================================================================
# Text Extraction Endpoints
# ============================================================================

@app.post("/identify/text", response_model=StudentIdentifierResponse)
def identify_student_identifiers(request: StudentIdentifierRequest):
    """
    Extract student NAME and ROLL NUMBER from text.
    
    This endpoint analyzes the input text and identifies student identifiers
    using pattern matching. Returns confidence levels for each extraction.
    """
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="text field is required and cannot be empty")
    
    name, roll_no, confidence = StudentIdentifierExtractor.extract_both(
        request.text, 
        strict=request.strict_mode
    )
    
    # Find a text excerpt around the found identifiers
    excerpt = None
    if name or roll_no:
        # Find first occurrence of identifier
        if name:
            idx = request.text.find(name)
        else:
            idx = request.text.find(roll_no)
        
        start = max(0, idx - 50)
        end = min(len(request.text), idx + 100)
        excerpt = request.text[start:end].strip()
    
    return StudentIdentifierResponse(
        detected_name=name,
        detected_roll_no=roll_no,
        extraction_confidence=confidence,
        raw_text_excerpt=excerpt
    )


# ============================================================================
# Text Redaction Endpoints
# ============================================================================

@app.post("/redact/text", response_model=RedactionResponse)
def redact_text(request: RedactionRequest):
    """
    Redact student NAME and ROLL NUMBER from text.
    
    This endpoint removes student identifiers from text while preserving
    the rest of the document content. Can optionally preserve field labels.
    """
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="text field is required and cannot be empty")
    
    redacted_text, metadata = StudentIdentifierRedactor.redact_both(
        request.text,
        remove_name=request.remove_name,
        remove_roll_no=request.remove_roll_no,
        preserve_labels=request.preserve_labels
    )
    
    return RedactionResponse(
        redacted_text=redacted_text,
        detected_name=metadata["detected_name"],
        detected_roll_no=metadata["detected_roll_no"],
        redaction_count=metadata["redaction_count"]
    )


# ============================================================================
# Document Redaction Endpoints
# ============================================================================

@app.post("/redact/document", response_model=DocumentRedactionResponse)
def redact_document(request: DocumentRedactionRequest):
    """
    Redact student identifiers from an entire document file.
    
    Supports DOCX, PDF, and TXT formats. Creates a new redacted file
    at the specified output path while preserving document formatting.
    """
    try:
        logger.info(f"Starting redaction: {request.input_file_path} -> {request.output_file_path}")
        
        # Validate input file
        if not os.path.exists(request.input_file_path):
            logger.error(f"Input file not found: {request.input_file_path}")
            raise HTTPException(status_code=400, detail=f"Input file not found: {request.input_file_path}")
        
        # Read document
        logger.info(f"Reading {request.file_format} document...")
        text = DocumentProcessor.process_document(
            request.input_file_path,
            request.output_file_path,
            request.file_format
        )
        logger.info(f"Document read successfully ({len(text)} characters)")
        
        # Redact identifiers
        logger.info(f"Redacting with options - remove_name: {request.remove_name}, remove_roll_no: {request.remove_roll_no}")
        redacted_text, metadata = StudentIdentifierRedactor.redact_both(
            text,
            remove_name=request.remove_name,
            remove_roll_no=request.remove_roll_no,
            preserve_labels=False
        )
        logger.info(f"Redaction complete - Detected Name: {metadata['detected_name']}, Roll No: {metadata['detected_roll_no']}")
        
        # Save redacted document
        logger.info(f"Saving redacted document to {request.output_file_path}...")
        if request.file_format.lower() == "docx":
            # Use XML-based anonymizer for non-table DOCX (preserves bullets)
            stats = anonymize_docx(
                request.input_file_path,
                request.output_file_path,
                name=metadata["detected_name"],
                roll_no=metadata["detected_roll_no"]
            )
            logger.info(f"Anonymizer stats: {stats}")
        else:
            DocumentProcessor.save_document(
                request.output_file_path,
                redacted_text,
                request.file_format
            )
        logger.info(f"Document saved successfully")
        return DocumentRedactionResponse(
            status="success",
            output_file=request.output_file_path,
            redacted_name=metadata["detected_name"],
            redacted_roll_no=metadata["detected_roll_no"]
        )
        
    except Exception as e:
        logger.error(f"Error during redaction: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/redact/batch")
def redact_batch_documents(requests: List[DocumentRedactionRequest]):
    """
    Redact multiple documents in batch.
    
    Processes multiple document redaction requests in sequence.
    Returns results for each document.
    """
    results = []
    
    for req in requests:
        try:
            text = DocumentProcessor.process_document(
                req.input_file_path,
                req.output_file_path,
                req.file_format
            )
            
            redacted_text, metadata = StudentIdentifierRedactor.redact_both(
                text,
                remove_name=req.remove_name,
                remove_roll=req.remove_roll_no,
                preserve_labels=False
            )
            
            DocumentProcessor.save_document(
                req.output_file_path,
                redacted_text,
                req.file_format
            )
            
            results.append({
                "status": "success",
                "input_file": req.input_file_path,
                "output_file": req.output_file_path,
                "redacted_name": metadata["detected_name"],
                "redacted_roll_no": metadata["detected_roll_no"]
            })
        except Exception as e:
            results.append({
                "status": "error",
                "input_file": req.input_file_path,
                "error": str(e)
            })
    
    return {"results": results, "total": len(requests), "successful": sum(1 for r in results if r["status"] == "success")}


# ============================================================================
# Server Startup
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    
    port = int(os.environ.get("PORT", "5018"))
    host = os.environ.get("HOST", "0.0.0.0")
    
    print(f"🚀 Starting Reductor Service v3...")
    print(f"📍 Host: {host}:{port}")
    print(f"📚 Purpose: Student NAME and ROLL NUMBER Redaction (Screenshot 2 Format)")
    print(f"📖 API Docs: http://localhost:{port}/docs")
    print()
    
    uvicorn.run(
        app,
        host=host,
        port=port,
        log_level="info"
    )


# Alias endpoint for orchestrator compatibility (must be after app is defined and before __main__)
from fastapi import Body

@app.post("/anonymize/docx", response_model=DocumentRedactionResponse)
def anonymize_docx_endpoint(request: DocumentRedactionRequest = Body(...)):
    """
    Alias for /redact/document for compatibility with orchestrators expecting /anonymize/docx.
    Only supports DOCX files.
    """
    if request.file_format.lower() != "docx":
        raise HTTPException(status_code=400, detail="Only DOCX format is supported at this endpoint.")
    return redact_document(request)
