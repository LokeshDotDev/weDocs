"""
Test script to verify bullet/list preservation in DOCX modification
"""
from docx import Document
import os

def test_bullet_preservation():
    """Test if modifying a DOCX preserves bullets"""
    
    # Create a test document with bullets
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('NAME: TEST STUDENT')
    doc.add_paragraph('ROLL NUMBER: 12345')
    doc.add_heading('List Section', 1)
    doc.add_paragraph('First bullet item', style='List Bullet')
    doc.add_paragraph('Second bullet item', style='List Bullet')
    doc.add_paragraph('Third bullet item', style='List Bullet')
    
    test_file = '/tmp/test_bullets_original.docx'
    output_file = '/tmp/test_bullets_modified.docx'
    
    doc.save(test_file)
    print(f"✓ Created test file: {test_file}")
    
    # Now modify it using the same approach as our service
    doc2 = Document(test_file)
    
    for para_idx, para in enumerate(doc2.paragraphs):
        para_text = para.text
        if 'NAME:' in para_text:
            print(f"Found NAME line at para {para_idx}: {para_text}")
            # Remove text after colon
            for run in para.runs:
                if 'TEST STUDENT' in run.text:
                    run.text = run.text.replace('TEST STUDENT', '')
                    print(f"  Modified run, new text: '{run.text}'")
    
    doc2.save(output_file)
    print(f"✓ Saved modified file: {output_file}")
    
    # Check if bullets are preserved
    doc3 = Document(output_file)
    print("\nChecking modified document:")
    for para_idx, para in enumerate(doc3.paragraphs):
        if para.style.name.startswith('List'):
            print(f"  Para {para_idx}: style='{para.style.name}', text='{para.text}'")
    
    print("\n✓ Test complete - check the files manually to verify bullets")
    print(f"  Original: {test_file}")
    print(f"  Modified: {output_file}")

if __name__ == "__main__":
    test_bullet_preservation()
