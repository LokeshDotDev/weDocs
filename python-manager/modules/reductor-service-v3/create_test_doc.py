#!/usr/bin/env python3
"""
Create a test DOCX with NAME and ROLL NUMBER for testing
"""

from docx import Document
import tempfile
import os

# Create test document
doc = Document()

# Add paragraphs with student info
doc.add_paragraph("NAME: SHANMUGAPRIYA SIVAKUMAR")
doc.add_paragraph("ROLL NUMBER: 251450500104")
doc.add_paragraph("PROGRAM: MASTER OF BUSINESS ADMINISTRATION(MBA)")
doc.add_paragraph("SEMESTER:1")
doc.add_paragraph("COURSE CODE & NAME: DMBA114- BUSINESS COMMUNICATION")

doc.add_paragraph("")
doc.add_paragraph("Assignment Content")
doc.add_paragraph("This is the assignment submitted by SHANMUGAPRIYA SIVAKUMAR (Roll: 251450500104)")
doc.add_paragraph("Some course materials and content here...")

# Save to temp location
output_path = "/tmp/test_business_communication.docx"
doc.save(output_path)
print(f"✓ Test document created: {output_path}")
print(f"  File size: {os.path.getsize(output_path)} bytes")
